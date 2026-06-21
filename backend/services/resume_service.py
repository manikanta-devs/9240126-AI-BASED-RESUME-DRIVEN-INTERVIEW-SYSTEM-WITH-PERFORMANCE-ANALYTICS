import logging
import re
import os

logger = logging.getLogger(__name__)


class ResumeService:
    """Service for extracting and analyzing resume content"""

    SKILL_KEYWORDS = {
        "languages": [
            "python",
            "javascript",
            "java",
            "c++",
            "c#",
            "ruby",
            "go",
            "rust",
            "typescript",
            "kotlin",
            "swift",
            "php",
            "scala",
            "r",
            "matlab",
            "perl",
            "bash",
            "shell",
            "dart",
            "elixir",
            "haskell",
        ],
        "frameworks": [
            "react",
            "angular",
            "vue",
            "django",
            "flask",
            "fastapi",
            "spring",
            "express",
            "node",
            "nodejs",
            "next.js",
            "nuxt",
            "svelte",
            "laravel",
            "rails",
            "asp.net",
            "tensorflow",
            "pytorch",
            "keras",
            "scikit-learn",
            "pandas",
            "numpy",
            "fastapi",
            "graphql",
            "rest",
            "redux",
            "tailwind",
        ],
        "databases": [
            "mysql",
            "postgresql",
            "mongodb",
            "redis",
            "sqlite",
            "oracle",
            "cassandra",
            "elasticsearch",
            "dynamodb",
            "firebase",
            "supabase",
            "neo4j",
            "influxdb",
            "mariadb",
            "sql server",
        ],
        "cloud_devops": [
            "aws",
            "azure",
            "gcp",
            "docker",
            "kubernetes",
            "jenkins",
            "git",
            "github",
            "gitlab",
            "ci/cd",
            "terraform",
            "ansible",
            "linux",
            "nginx",
            "apache",
            "heroku",
            "vercel",
            "netlify",
            "cloudflare",
        ],
        "tools": [
            "jira",
            "confluence",
            "slack",
            "figma",
            "postman",
            "swagger",
            "vs code",
            "intellij",
            "xcode",
            "android studio",
            "jupyter",
            "tableau",
            "power bi",
            "excel",
            "git",
            "webpack",
            "vite",
        ],
    }

    EDUCATION_KEYWORDS = [
        "bachelor",
        "master",
        "phd",
        "doctorate",
        "b.sc",
        "m.sc",
        "b.tech",
        "m.tech",
        "mba",
        "b.e",
        "m.e",
        "degree",
        "diploma",
        "certification",
        "certificate",
        "university",
        "college",
        "institute",
        "school",
        "computer science",
        "information technology",
        "engineering",
        "graduate",
    ]

    EXPERIENCE_PATTERNS = [
        r"(\d+)\+?\s*years?(?:\s+(?:of\s+)?experience)?",
        r"(\d+)\+?\s*years?\s+(?:of\s+)?experience",
        r"experience\s+(?:of\s+)?(\d+)\+?\s*years?",
        r"(\d+)\+?\s*yrs?\s+(?:of\s+)?experience",
    ]

    def __init__(self):
        self.nlp = self._load_spacy()

    def _load_spacy(self):
        """Load spaCy model with fallback"""
        try:
            import spacy

            try:
                return spacy.load("en_core_web_sm")
            except OSError:
                logger.warning(
                    "spaCy model not found. Run: python -m spacy download en_core_web_sm"
                )
                return None
        except Exception as exc:
            logger.warning(
                "spaCy unavailable, falling back to regex extraction: %s", exc
            )
            return None

    def extract_text_from_pdf(self, filepath: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2

            text = []
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return "\n".join(text)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to extract PDF text: {e}")

    def extract_text_from_docx(self, filepath: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document

            doc = Document(filepath)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            # Extract from tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            return "\n".join(text)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to extract DOCX text: {e}")

    def extract_text_from_txt(self, filepath: str) -> str:
        """Extract text from TXT file"""
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def analyze_resume(self, filepath: str) -> dict:
        """Main method to analyze a resume file"""
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".pdf":
            text = self.extract_text_from_pdf(filepath)
        elif ext in [".docx", ".doc"]:
            text = self.extract_text_from_docx(filepath)
        elif ext == ".txt":
            text = self.extract_text_from_txt(filepath)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        return self.analyze_text(text)

    def analyze_text(self, text: str) -> dict:
        """Analyze resume text and extract structured information"""
        text_lower = text.lower()

        skills = self._extract_skills(text_lower)
        education = self._extract_education(text, text_lower)
        experience = self._extract_experience(text, text_lower)
        contact = self._extract_contact(text)
        entities = self._extract_entities(text)
        score = self._calculate_resume_score(skills, education, experience, text)

        return {
            "raw_text_length": len(text),
            "skills": skills,
            "education": education,
            "experience": experience,
            "contact": contact,
            "entities": entities,
            "resume_score": score,
            "summary": self._generate_summary(skills, education, experience),
        }

    def _extract_skills(self, text_lower: str) -> dict:
        """Extract skills by category"""
        found = {}
        all_skills = []
        for category, keywords in self.SKILL_KEYWORDS.items():
            matched = []
            for kw in keywords:
                if kw in text_lower:
                    matched.append(kw.title() if len(kw) > 3 else kw.upper())
            if matched:
                found[category] = list(set(matched))
                all_skills.extend(matched)
        found["all"] = list(set(all_skills))
        return found

    def _extract_education(self, text: str, text_lower: str) -> list:
        """Extract education information"""
        education = []
        lines = text.split("\n")
        for line in lines:
            line_lower = line.lower()
            if any(kw in line_lower for kw in self.EDUCATION_KEYWORDS):
                clean = line.strip()
                if len(clean) > 10 and clean not in education:
                    education.append(clean)
        return education[:5]  # Limit to 5 entries

    def _extract_experience(self, text: str, text_lower: str) -> dict:
        """Extract work experience information"""
        years = None
        for pattern in self.EXPERIENCE_PATTERNS:
            match = re.search(pattern, text_lower)
            if match:
                years = int(match.group(1))
                break

        job_titles = []
        title_keywords = [
            "engineer",
            "developer",
            "manager",
            "analyst",
            "designer",
            "architect",
            "lead",
            "senior",
            "junior",
            "intern",
            "consultant",
            "specialist",
            "director",
            "head of",
            "vp",
            "cto",
            "ceo",
        ]
        lines = text.split("\n")
        for line in lines:
            line_lower = line.lower()
            if any(kw in line_lower for kw in title_keywords):
                clean = line.strip()
                if 5 < len(clean) < 80 and clean not in job_titles:
                    job_titles.append(clean)

        level = "entry"
        if years:
            if years >= 8:
                level = "senior"
            elif years >= 4:
                level = "mid"
            elif years >= 1:
                level = "junior"

        return {"years": years, "level": level, "titles": job_titles[:5]}

    def _extract_contact(self, text: str) -> dict:
        """Extract contact information"""
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        phone_pattern = r"[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}"
        linkedin_pattern = r"linkedin\.com/in/[\w-]+"
        github_pattern = r"github\.com/[\w-]+"

        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        linkedin = re.findall(linkedin_pattern, text.lower())
        github = re.findall(github_pattern, text.lower())

        return {
            "email": emails[0] if emails else None,
            "phone": phones[0] if phones else None,
            "linkedin": linkedin[0] if linkedin else None,
            "github": github[0] if github else None,
        }

    def _extract_entities(self, text: str) -> dict:
        """Use spaCy to extract named entities"""
        if not self.nlp:
            return {"organizations": [], "locations": [], "persons": []}

        try:
            doc = self.nlp(text[:5000])  # Limit for performance
            entities = {"organizations": [], "locations": [], "persons": []}

            for ent in doc.ents:
                if ent.label_ == "ORG" and ent.text not in entities["organizations"]:
                    entities["organizations"].append(ent.text)
                elif (
                    ent.label_ in ["GPE", "LOC"]
                    and ent.text not in entities["locations"]
                ):
                    entities["locations"].append(ent.text)
                elif ent.label_ == "PERSON" and ent.text not in entities["persons"]:
                    entities["persons"].append(ent.text)

            return {k: v[:10] for k, v in entities.items()}
        except Exception as e:
            logger.error(f"spaCy entity extraction error: {e}")
            return {"organizations": [], "locations": [], "persons": []}

    def _calculate_resume_score(
        self, skills: dict, education: list, experience: dict, text: str
    ) -> dict:
        """Calculate a resume quality score"""
        scores = {}
        # Skills score (0-25)
        skill_count = len(skills.get("all", []))
        scores["skills"] = min(25, skill_count * 2)

        # Education score (0-25)
        scores["education"] = min(25, len(education) * 8)

        # Experience score (0-25)
        exp_years = experience.get("years", 0) or 0
        scores["experience"] = min(25, exp_years * 4)

        # Completeness score (0-25)
        completeness = 0
        if len(text) > 500:
            completeness += 5
        if len(text) > 1000:
            completeness += 5
        if "email" in text.lower() or "@" in text:
            completeness += 5
        if any(w in text.lower() for w in ["project", "achievement", "award"]):
            completeness += 5
        if len(education) > 0:
            completeness += 5
        scores["completeness"] = completeness

        total = sum(scores.values())
        grades = {90: "A+", 80: "A", 70: "B+", 60: "B", 50: "C+", 0: "C"}
        grade = next(
            g
            for threshold, g in sorted(grades.items(), reverse=True)
            if total >= threshold
        )

        return {
            "total": total,
            "max": 100,
            "grade": grade,
            "breakdown": scores,
            "percentage": total,
        }

    def _generate_summary(self, skills: dict, education: list, experience: dict) -> str:
        """Generate a brief resume summary"""
        parts = []
        exp_years = experience.get("years")
        level = experience.get("level", "entry")

        if exp_years:
            parts.append(f"{exp_years} years of experience")

        skill_count = len(skills.get("all", []))
        if skill_count > 0:
            top_skills = skills.get("all", [])[:3]
            parts.append(f"skilled in {', '.join(top_skills)}")

        if education:
            parts.append("with formal education background")

        if parts:
            return f"{level.title()}-level professional {', '.join(parts)}."
        return "Resume processed successfully."

    def compare_resume_to_job(self, resume_data: dict, job_description: str) -> dict:
        """Compare analyzed resume data against a target job description."""
        if not resume_data:
            raise ValueError("Resume data is required for job matching")

        if not job_description or len(job_description.strip()) < 30:
            raise ValueError("Job description is required for job matching")

        job_analysis = self.analyze_text(job_description)

        resume_skill_map = self._skill_map(resume_data.get("skills", {}).get("all", []))
        job_skill_map = self._skill_map(job_analysis.get("skills", {}).get("all", []))

        resume_skill_keys = set(resume_skill_map.keys())
        job_skill_keys = set(job_skill_map.keys())

        matched_keys = sorted(resume_skill_keys & job_skill_keys)
        missing_keys = sorted(job_skill_keys - resume_skill_keys)
        extra_keys = sorted(resume_skill_keys - job_skill_keys)

        matched_skills = [
            resume_skill_map.get(key) or job_skill_map.get(key) for key in matched_keys
        ]
        missing_skills = [job_skill_map[key] for key in missing_keys]
        extra_skills = [resume_skill_map[key] for key in extra_keys]

        resume_experience = resume_data.get("experience", {}) or {}
        job_experience = job_analysis.get("experience", {}) or {}
        resume_years = int(resume_experience.get("years") or 0)
        required_years = int(job_experience.get("years") or 0)

        skill_coverage = (
            len(matched_skills) / len(job_skill_keys) if job_skill_keys else 0.65
        )
        experience_alignment = 1.0
        if required_years > 0:
            experience_alignment = (
                min(1.0, resume_years / required_years) if resume_years else 0.0
            )

        resume_completion = (resume_data.get("resume_score", {}) or {}).get(
            "percentage", 0
        ) / 100
        education_alignment = 1.0 if resume_data.get("education") else 0.6

        total_score = round(
            (skill_coverage * 58)
            + (experience_alignment * 20)
            + (education_alignment * 10)
            + (min(1.0, resume_completion) * 12)
        )
        total_score = max(0, min(100, total_score))

        grade = self._grade_from_score(total_score)
        readiness = self._readiness_label(total_score)

        strengths = []
        if matched_skills:
            strengths.append(f"Matches {len(matched_skills)} job keywords")
        if required_years and resume_years >= required_years:
            strengths.append("Experience level meets or exceeds the job signal")
        if resume_data.get("education"):
            strengths.append("Education section is present")
        if resume_data.get("summary"):
            strengths.append("Resume already has a concise summary")

        recommendations = []
        if missing_skills:
            recommendations.append(
                f"Add or emphasize these keywords: {', '.join(missing_skills[:5])}."
            )
        if required_years and resume_years < required_years:
            recommendations.append(
                f"Highlight projects that show at least {required_years - resume_years} more year(s) of relevant impact."
            )
        if not resume_data.get("contact", {}).get("email"):
            recommendations.append(
                "Make sure your contact details are visible and complete."
            )
        if len(recommendations) < 3:
            recommendations.append(
                "Tailor your professional summary to mirror the language in the job description."
            )
            recommendations.append(
                "Quantify outcomes in bullets with metrics, scale, or impact."
            )

        priority_actions = [
            "Rewrite the top summary line with the role title and 2-3 matching keywords.",
            "Add one bullet for each missing priority skill using a concrete project example.",
            "Prepare interview stories around the matched skills and the strongest project wins.",
        ]

        return {
            "match_score": total_score,
            "match_grade": grade,
            "readiness": readiness,
            "matched_skills": matched_skills[:12],
            "missing_skills": missing_skills[:12],
            "extra_skills": extra_skills[:12],
            "resume_years": resume_years,
            "required_years": required_years,
            "skill_coverage": round(skill_coverage * 100, 1),
            "experience_alignment": round(experience_alignment * 100, 1),
            "education_alignment": round(education_alignment * 100, 1),
            "strengths": strengths[:4],
            "recommendations": recommendations[:4],
            "priority_actions": priority_actions,
            "job_keywords": list(job_skill_map.values())[:12],
            "summary": self._build_match_summary(
                total_score,
                matched_skills,
                missing_skills,
                required_years,
                resume_years,
            ),
        }

    def _skill_map(self, skills: list) -> dict:
        """Normalize skill labels for matching while keeping display text."""
        mapped = {}
        for skill in skills or []:
            if not skill:
                continue
            display = skill.strip()
            normalized = display.lower()
            mapped[normalized] = display
        return mapped

    def _grade_from_score(self, score: int) -> str:
        """Convert a numeric score to a letter grade."""
        grade_map = [(90, "A+"), (80, "A"), (70, "B+"), (60, "B"), (50, "C+"), (0, "C")]
        return next(grade for threshold, grade in grade_map if score >= threshold)

    def _readiness_label(self, score: int) -> str:
        """Translate the fit score into a hiring-readiness label."""
        if score >= 85:
            return "interview-ready"
        if score >= 70:
            return "strong-match"
        if score >= 55:
            return "build-gap"
        return "foundation-needed"

    def _build_match_summary(
        self,
        score: int,
        matched_skills: list,
        missing_skills: list,
        required_years: int,
        resume_years: int,
    ) -> str:
        """Generate a short executive summary for the match report."""
        if score >= 85:
            prefix = "Excellent alignment"
        elif score >= 70:
            prefix = "Strong alignment"
        elif score >= 55:
            prefix = "Moderate alignment"
        else:
            prefix = "Low alignment"

        skill_text = f"{len(matched_skills)} matched skills"
        gap_text = f"{len(missing_skills)} keyword gaps"
        if required_years:
            exp_text = f"{resume_years}/{required_years} years of experience signal"
        else:
            exp_text = f"{resume_years} years of experience signal"

        return f"{prefix}: {skill_text}, {gap_text}, {exp_text}."
